"""
Frontier AI Company Dashboard — visualizations and data layer.
"""

from __future__ import annotations
import io
import json
from datetime import datetime
from pathlib import Path

import pandas as pd
import plotly.express as px
import plotly.graph_objects as go
import streamlit as st
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

DATA_FILE = Path(__file__).parent.parent / "data" / "companies.json"

REGION_COLORS = {"US": "#1f77b4", "PRC": "#d62728", "EU": "#2ca02c"}
REGION_LABELS = {"US": "🇺🇸 United States", "PRC": "🇨🇳 China", "EU": "🇪🇺 Europe"}

RISK_LABELS = {
    "compute_dependency": "Compute Dependency",
    "geopolitical_exposure": "Geopolitical Exposure",
    "concentration_risk": "Concentration Risk",
    "regulatory_risk": "Regulatory Risk",
    "supply_chain_risk": "Supply Chain Risk",
}


# ── Data loading ───────────────────────────────────────────────────────────────

@st.cache_data(ttl=60)
def load_companies() -> list[dict]:
    return json.loads(DATA_FILE.read_text(encoding="utf-8"))


def to_dataframe(companies: list[dict]) -> pd.DataFrame:
    rows = []
    for c in companies:
        market_cap = c.get("parent_market_cap_bn") or c.get("valuation_bn")
        rows.append({
            "Company": c["name"],
            "Region": c["region"],
            "Founded": c["founded"],
            "Status": c["status"],
            "Funding / Val ($B)": market_cap,
            "Employees": c.get("employees_approx"),
            "Latest Model": c.get("latest_model", "—"),
            "Open Source": "✅" if c.get("open_source") else "❌",
            "Compute Partner": c.get("compute_partner", "—"),
            "Own Chip": "✅" if c.get("proprietary_chip") else "❌",
            "Est. GPUs (K)": c.get("estimated_gpus_k"),
            "Capability Score": c.get("capability_score"),
            "Gov Ties": c.get("government_ties", "—"),
            "Export Control Risk": c.get("export_control_risk", "—"),
            "id": c["id"],
        })
    return pd.DataFrame(rows)


def avg_risk(c: dict) -> float:
    risks = c.get("risks", {})
    vals = list(risks.values())
    return round(sum(vals) / len(vals), 1) if vals else 0


# ── Overview table ─────────────────────────────────────────────────────────────

def render_overview(companies: list[dict], regions: list[str]):
    filtered = [c for c in companies if c["region"] in regions]
    if not filtered:
        st.info("No companies match the selected filters.")
        return

    df = to_dataframe(filtered).drop(columns=["id"])

    # Color-code by region
    def row_color(row):
        clr = {"US": "#e8f0fe", "PRC": "#fde8e8", "EU": "#e8fde8"}.get(row["Region"], "")
        return [f"background-color: {clr}" for _ in row]

    st.dataframe(
        df.style.apply(row_color, axis=1),
        use_container_width=True,
        height=420,
        hide_index=True,
    )

    # Summary metrics row
    st.divider()
    mc1, mc2, mc3, mc4 = st.columns(4)
    mc1.metric("Companies tracked", len(filtered))
    mc2.metric("US labs", sum(1 for c in filtered if c["region"] == "US"))
    mc3.metric("PRC labs", sum(1 for c in filtered if c["region"] == "PRC"))
    mc4.metric("EU labs", sum(1 for c in filtered if c["region"] == "EU"))


# ── Global map ─────────────────────────────────────────────────────────────────

def render_map(companies: list[dict], regions: list[str]):
    filtered = [c for c in companies if c["region"] in regions]

    map_data = []
    for c in filtered:
        models_str = ", ".join(c.get("key_models", [])[:3])
        map_data.append({
            "name": c["name"],
            "region": c["region"],
            "lat": c["hq_lat"],
            "lon": c["hq_lon"],
            "city": c["hq_city"],
            "capability": c.get("capability_score", 5),
            "gpus_k": c.get("estimated_gpus_k", 10),
            "hover": (
                f"<b>{c['name']}</b><br>"
                f"📍 {c['hq_city']}<br>"
                f"🤖 {c.get('latest_model','—')}<br>"
                f"💻 {c.get('estimated_gpus_k','?')}K GPUs (est.)<br>"
                f"⚡ Capability: {c.get('capability_score','?')}/10"
            ),
        })

    df = pd.DataFrame(map_data)

    fig = px.scatter_geo(
        df,
        lat="lat",
        lon="lon",
        color="region",
        size="gpus_k",
        hover_name="name",
        custom_data=["hover"],
        color_discrete_map=REGION_COLORS,
        size_max=50,
        projection="natural earth",
        title="Frontier AI Lab HQ Locations (bubble = estimated GPU scale)",
    )
    fig.update_traces(
        hovertemplate="%{customdata[0]}<extra></extra>",
        marker=dict(opacity=0.8, line=dict(width=1, color="white")),
    )
    fig.update_layout(
        height=500,
        margin=dict(l=0, r=0, t=40, b=0),
        legend_title="Region",
        geo=dict(
            showland=True,
            landcolor="rgb(243,243,243)",
            showocean=True,
            oceancolor="rgb(220,235,255)",
            showcountries=True,
            countrycolor="rgb(200,200,200)",
        ),
    )
    st.plotly_chart(fig, use_container_width=True)
    st.caption("Bubble size = estimated GPU count (K). HQ locations shown; data center footprints not depicted.")


# ── Timeline ───────────────────────────────────────────────────────────────────

def render_timeline(companies: list[dict], regions: list[str]):
    filtered = [c for c in companies if c["region"] in regions]

    events = []
    for c in filtered:
        for m in c.get("milestones", []):
            events.append({
                "Company": c["name"],
                "Region": c["region"],
                "Date": m["date"],
                "Event": m["event"],
            })

    if not events:
        st.info("No milestone data available.")
        return

    df = pd.DataFrame(events)
    df["Date"] = pd.to_datetime(df["Date"])
    df = df.sort_values("Date")

    fig = px.scatter(
        df,
        x="Date",
        y="Company",
        color="Region",
        hover_data={"Event": True, "Date": True, "Region": False},
        color_discrete_map=REGION_COLORS,
        title="Frontier AI Milestone Timeline",
        size_max=12,
    )
    fig.update_traces(marker=dict(size=10, symbol="diamond"))
    fig.update_layout(
        height=500,
        xaxis_title="",
        yaxis_title="",
        legend_title="Region",
        hovermode="closest",
        margin=dict(l=10, r=10, t=40, b=10),
    )
    st.plotly_chart(fig, use_container_width=True)

    # Expandable raw list
    with st.expander("View all milestones"):
        for _, row in df.sort_values("Date", ascending=False).iterrows():
            color = {"US": "🔵", "PRC": "🔴", "EU": "🟢"}.get(row["Region"], "⚪")
            st.markdown(
                f"{color} **{row['Date'].strftime('%b %Y')}** · {row['Company']} — {row['Event']}"
            )


# ── Risk matrix ────────────────────────────────────────────────────────────────

def render_risk_matrix(companies: list[dict], regions: list[str], risk_x: str, risk_y: str):
    filtered = [c for c in companies if c["region"] in regions]

    rows = []
    for c in filtered:
        risks = c.get("risks", {})
        models_str = " · ".join(c.get("key_models", [])[:2])
        rows.append({
            "Company": c["name"],
            "Region": c["region"],
            risk_x: risks.get(risk_x, 5),
            risk_y: risks.get(risk_y, 5),
            "Capability": c.get("capability_score", 5),
            "GPUs (K)": c.get("estimated_gpus_k", 10),
            "Models": models_str,
        })

    df = pd.DataFrame(rows)

    fig = px.scatter(
        df,
        x=risk_x,
        y=risk_y,
        size="GPUs (K)",
        color="Region",
        text="Company",
        hover_data={"Models": True, "Capability": True, "GPUs (K)": True},
        color_discrete_map=REGION_COLORS,
        size_max=55,
        title=f"Risk Matrix: {RISK_LABELS[risk_x]} vs {RISK_LABELS[risk_y]}",
    )
    fig.update_traces(
        textposition="top center",
        marker=dict(opacity=0.75, line=dict(width=1, color="white")),
    )

    # Quadrant shading
    fig.add_shape(type="rect", x0=5, y0=5, x1=10.5, y1=10.5,
                  fillcolor="rgba(255,0,0,0.05)", line=dict(width=0))
    fig.add_annotation(x=7.5, y=10, text="HIGH RISK ZONE", showarrow=False,
                       font=dict(color="rgba(200,0,0,0.4)", size=11))

    fig.update_layout(
        height=520,
        xaxis=dict(title=RISK_LABELS[risk_x], range=[0, 11], dtick=1),
        yaxis=dict(title=RISK_LABELS[risk_y], range=[0, 11], dtick=1),
        legend_title="Region",
        margin=dict(l=10, r=10, t=50, b=10),
    )
    st.plotly_chart(fig, use_container_width=True)
    st.caption("Bubble size = estimated GPU scale. Scores are analyst estimates (1=low risk, 10=high risk).")


# ── Capability comparison bar chart ───────────────────────────────────────────

def render_capability_chart(companies: list[dict], regions: list[str]):
    filtered = sorted(
        [c for c in companies if c["region"] in regions],
        key=lambda c: c.get("capability_score", 0),
        reverse=True,
    )

    df = pd.DataFrame([{
        "Company": c["name"],
        "Region": c["region"],
        "Capability Score": c.get("capability_score", 0),
        "GPU Scale (K)": c.get("estimated_gpus_k", 0),
    } for c in filtered])

    fig = px.bar(
        df,
        x="Company",
        y="Capability Score",
        color="Region",
        color_discrete_map=REGION_COLORS,
        title="Frontier Capability Scores by Company",
        text="Capability Score",
    )
    fig.update_layout(
        height=380,
        yaxis=dict(range=[0, 11]),
        xaxis_title="",
        yaxis_title="Capability Score (analyst estimate, /10)",
        showlegend=True,
        margin=dict(t=50, b=10),
    )
    fig.update_traces(textposition="outside")
    st.plotly_chart(fig, use_container_width=True)


# ── Company deep dive ──────────────────────────────────────────────────────────

def render_deep_dive(company: dict):
    st.subheader(f"{company['name']}")

    region_flag = {"US": "🇺🇸", "PRC": "🇨🇳", "EU": "🇪🇺"}.get(company["region"], "🌐")
    c1, c2, c3, c4 = st.columns(4)

    val = company.get("parent_market_cap_bn") or company.get("valuation_bn")
    val_label = "Market Cap" if company.get("parent_market_cap_bn") else "Valuation"
    c1.metric("Region", f"{region_flag} {company['region']}")
    c2.metric("Founded", company.get("founded", "—"))
    c3.metric(val_label, f"${val}B" if val else "Private / N/A")
    c4.metric("Capability Score", f"{company.get('capability_score','—')} / 10")

    c5, c6, c7, c8 = st.columns(4)
    c5.metric("Latest Model", company.get("latest_model", "—"))
    c6.metric("Est. GPUs", f"{company.get('estimated_gpus_k','?')}K")
    c7.metric("Open Source", "Yes" if company.get("open_source") else "No")
    c8.metric("Export Control Risk", company.get("export_control_risk", "—"))

    st.divider()
    col_left, col_right = st.columns([3, 2])

    with col_left:
        st.markdown("**Strategic Notes**")
        st.markdown(company.get("strategic_notes", "—"))

        st.markdown("**Key Models**")
        for m in company.get("key_models", []):
            st.markdown(f"• {m}")

    with col_right:
        st.markdown("**Risk Profile**")
        risks = company.get("risks", {})
        risk_df = pd.DataFrame([
            {"Risk Dimension": RISK_LABELS.get(k, k), "Score": v}
            for k, v in risks.items()
        ])
        fig = px.bar(
            risk_df,
            x="Score",
            y="Risk Dimension",
            orientation="h",
            color="Score",
            color_continuous_scale=["#2ecc71", "#f39c12", "#e74c3c"],
            range_color=[1, 10],
            range_x=[0, 10],
        )
        fig.update_layout(
            height=250,
            showlegend=False,
            coloraxis_showscale=False,
            margin=dict(l=5, r=5, t=5, b=5),
            xaxis_title="Risk Score (1=low, 10=high)",
            yaxis_title="",
        )
        st.plotly_chart(fig, use_container_width=True)

    st.divider()
    st.markdown("**Key Milestones**")
    for m in sorted(company.get("milestones", []), key=lambda x: x["date"], reverse=True):
        date_str = datetime.strptime(m["date"], "%Y-%m-%d").strftime("%b %Y")
        st.markdown(f"**{date_str}** — {m['event']}")

    st.caption(f"Data last updated: {company.get('last_updated','—')} · All figures are public estimates.")


# ── Dashboard report export ────────────────────────────────────────────────────

def build_dashboard_report(companies: list[dict], regions: list[str]) -> bytes:
    filtered = [c for c in companies if c["region"] in regions]
    doc = Document()

    # Title
    title = doc.add_heading("Frontier AI Company Intelligence Report", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_p = doc.add_paragraph(datetime.now().strftime("%B %d, %Y"))
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_p.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    doc.add_paragraph()

    # Executive summary
    doc.add_heading("Executive Summary", level=1)
    us = [c for c in filtered if c["region"] == "US"]
    prc = [c for c in filtered if c["region"] == "PRC"]
    eu = [c for c in filtered if c["region"] == "EU"]
    summary = (
        f"This report covers {len(filtered)} frontier AI organizations across "
        f"{len(us)} US, {len(prc)} PRC, and {len(eu)} EU entities. "
        f"Data reflects publicly available information as of {datetime.now().strftime('%B %Y')}. "
        f"All capability scores, GPU estimates, and risk ratings are analyst estimates."
    )
    doc.add_paragraph(summary)

    # Company profiles
    doc.add_heading("Company Profiles", level=1)
    for c in filtered:
        doc.add_heading(c["name"], level=2)

        val = c.get("parent_market_cap_bn") or c.get("valuation_bn")
        val_label = "Market Cap" if c.get("parent_market_cap_bn") else "Valuation"

        info_lines = [
            f"Region: {c['region']}  |  Founded: {c['founded']}  |  Status: {c['status']}",
            f"{val_label}: {'$'+str(val)+'B' if val else 'N/A'}  |  Employees: ~{c.get('employees_approx','?'):,}",
            f"Latest Model: {c.get('latest_model','—')}  |  Open Source: {'Yes' if c.get('open_source') else 'No'}",
            f"Compute: {c.get('compute_partner','—')}  |  Primary Chip: {c.get('primary_chip','—')}",
            f"Est. GPUs: {c.get('estimated_gpus_k','?')}K  |  Capability Score: {c.get('capability_score','—')}/10",
            f"Export Control Risk: {c.get('export_control_risk','—')}  |  Gov Ties: {c.get('government_ties','—')}",
        ]
        for line in info_lines:
            p = doc.add_paragraph(line)
            p.paragraph_format.space_after = Pt(2)

        doc.add_paragraph()
        doc.add_paragraph(c.get("strategic_notes", ""))

        # Risk scores
        risks = c.get("risks", {})
        if risks:
            doc.add_paragraph()
            risk_p = doc.add_paragraph("Risk Scores: ")
            for k, v in risks.items():
                risk_p.add_run(f"{RISK_LABELS.get(k,k)}: {v}/10  ")

        doc.add_paragraph()

    # Risk comparison table
    doc.add_heading("Risk Comparison Matrix", level=1)
    risk_keys = list(RISK_LABELS.keys())
    table = doc.add_table(rows=1 + len(filtered), cols=1 + len(risk_keys))
    table.style = "Light List Accent 1"

    hdr = table.rows[0].cells
    hdr[0].text = "Company"
    for i, k in enumerate(risk_keys):
        hdr[i+1].text = RISK_LABELS[k]

    for r, c in enumerate(filtered, 1):
        row = table.rows[r].cells
        row[0].text = c["name"]
        for i, k in enumerate(risk_keys):
            row[i+1].text = str(c.get("risks", {}).get(k, "—"))

    doc.add_paragraph()
    doc.add_paragraph("Scores: 1 = low risk, 10 = high risk. All estimates.")

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# ── Scenario: Radar chart (before vs after) ────────────────────────────────────

def render_scenario_radar(company_result: dict):
    """Spider/radar chart comparing before and after risk profiles for one company."""
    import plotly.graph_objects as go

    risk_keys = list(RISK_LABELS.keys())
    labels = [RISK_LABELS[k] for k in risk_keys]

    before = [company_result["scenario_before"]["risks"].get(k, 5) for k in risk_keys]
    after  = [company_result["scenario_after"]["risks"].get(k, 5) for k in risk_keys]

    # Close the polygon
    labels_closed = labels + [labels[0]]
    before_closed = before + [before[0]]
    after_closed  = after  + [after[0]]

    fig = go.Figure()
    fig.add_trace(go.Scatterpolar(
        r=before_closed, theta=labels_closed,
        fill="toself", name="Before",
        line=dict(color="#1f77b4", dash="dash", width=2),
        fillcolor="rgba(31,119,180,0.08)",
    ))
    fig.add_trace(go.Scatterpolar(
        r=after_closed, theta=labels_closed,
        fill="toself", name="After",
        line=dict(color="#d62728", width=2.5),
        fillcolor="rgba(214,39,40,0.12)",
    ))

    cap_before = company_result["scenario_before"]["capability_score"]
    cap_after  = company_result["scenario_after"]["capability_score"]
    conf       = company_result.get("confidence", 1.0)

    fig.update_layout(
        polar=dict(radialaxis=dict(visible=True, range=[0, 10], dtick=2)),
        showlegend=True,
        title=dict(
            text=(
                f"{company_result['name']} — Risk Profile Shift<br>"
                f"<sup>Capability: {cap_before} → {cap_after}  |  "
                f"Confidence: {int(conf*100)}%</sup>"
            ),
            x=0.5,
        ),
        height=420,
        margin=dict(t=80, b=20, l=60, r=60),
    )
    st.plotly_chart(fig, use_container_width=True)


# ── Scenario: Risk matrix with movement arrows ─────────────────────────────────

def render_scenario_risk_matrix(
    results: list[dict],
    regions: list[str],
    risk_x: str,
    risk_y: str,
):
    """Risk matrix showing before positions (faded) and after positions (solid)
    with arrows indicating movement."""
    import plotly.graph_objects as go

    filtered = [r for r in results if r["region"] in regions]

    fig = go.Figure()

    for company in filtered:
        color = REGION_COLORS.get(company["region"], "#888")

        bx = company["scenario_before"]["risks"].get(risk_x, 5)
        by = company["scenario_before"]["risks"].get(risk_y, 5)
        ax = company["scenario_after"]["risks"].get(risk_x, 5)
        ay = company["scenario_after"]["risks"].get(risk_y, 5)

        # Before marker (faded)
        fig.add_trace(go.Scatter(
            x=[bx], y=[by],
            mode="markers",
            marker=dict(size=14, color=color, opacity=0.25,
                        line=dict(width=1, color=color)),
            showlegend=False,
            hoverinfo="skip",
        ))

        # After marker (solid)
        fig.add_trace(go.Scatter(
            x=[ax], y=[ay],
            mode="markers+text",
            marker=dict(size=16, color=color, opacity=0.85,
                        line=dict(width=1.5, color="white")),
            text=[company["name"]],
            textposition="top center",
            textfont=dict(size=10),
            name=company["name"],
            showlegend=False,
            hovertemplate=(
                f"<b>{company['name']}</b><br>"
                f"{RISK_LABELS[risk_x]}: {bx:.1f} → {ax:.1f}<br>"
                f"{RISK_LABELS[risk_y]}: {by:.1f} → {ay:.1f}<br>"
                f"Confidence: {int(company.get('confidence',1)*100)}%"
                "<extra></extra>"
            ),
        ))

        # Arrow if meaningful movement
        if abs(ax - bx) > 0.05 or abs(ay - by) > 0.05:
            fig.add_annotation(
                x=ax, y=ay, ax=bx, ay=by,
                xref="x", yref="y", axref="x", ayref="y",
                showarrow=True,
                arrowhead=3, arrowsize=1.2, arrowwidth=1.5,
                arrowcolor=color,
            )

    # Region legend dots
    for region, color in REGION_COLORS.items():
        if region in regions:
            fig.add_trace(go.Scatter(
                x=[None], y=[None], mode="markers",
                marker=dict(size=10, color=color),
                name=REGION_LABELS.get(region, region),
            ))

    # High-risk quadrant shading
    fig.add_shape(type="rect", x0=5, y0=5, x1=10.5, y1=10.5,
                  fillcolor="rgba(255,0,0,0.04)", line=dict(width=0))
    fig.add_annotation(x=7.5, y=10, text="HIGH RISK ZONE", showarrow=False,
                       font=dict(color="rgba(200,0,0,0.35)", size=10))

    fig.update_layout(
        title=f"Risk Matrix Movement: {RISK_LABELS[risk_x]} vs {RISK_LABELS[risk_y]}",
        xaxis=dict(title=RISK_LABELS[risk_x], range=[0, 11], dtick=1),
        yaxis=dict(title=RISK_LABELS[risk_y], range=[0, 11], dtick=1),
        height=520,
        margin=dict(l=10, r=10, t=50, b=10),
        showlegend=True,
        legend_title="Region",
    )
    st.plotly_chart(fig, use_container_width=True)
    st.caption("Faded dot = baseline position. Solid dot = projected position. Arrow = direction of change.")


# ── Scenario: Capability ranking change ────────────────────────────────────────

def render_scenario_capability_shift(results: list[dict], regions: list[str]):
    """Before/after capability bar chart side by side."""
    filtered = [r for r in results if r["region"] in regions]
    filtered_sorted = sorted(filtered, key=lambda r: r["scenario_after"]["capability_score"], reverse=True)

    rows = []
    for r in filtered_sorted:
        cap_b = r["scenario_before"]["capability_score"]
        cap_a = r["scenario_after"]["capability_score"]
        delta = round(cap_a - cap_b, 2)
        rows.append({
            "Company": r["name"],
            "Region": r["region"],
            "Before": cap_b,
            "After": cap_a,
            "Δ": delta,
            "Confidence": f"{int(r.get('confidence',1)*100)}%",
        })

    df = pd.DataFrame(rows)

    fig = go.Figure()
    for region, color in REGION_COLORS.items():
        mask = df["Region"] == region
        if not mask.any():
            continue
        sub = df[mask]
        fig.add_trace(go.Bar(
            name=f"{REGION_LABELS.get(region, region)} — Before",
            x=sub["Company"], y=sub["Before"],
            marker_color=color, opacity=0.35,
            offsetgroup=region,
        ))
        fig.add_trace(go.Bar(
            name=f"{REGION_LABELS.get(region, region)} — After",
            x=sub["Company"], y=sub["After"],
            marker_color=color, opacity=0.9,
            offsetgroup=region + "_after",
        ))

    fig.update_layout(
        barmode="group",
        title="Capability Score: Before vs After",
        yaxis=dict(title="Capability Score (/10)", range=[0, 11]),
        xaxis_title="",
        height=380,
        margin=dict(t=50, b=10),
    )
    st.plotly_chart(fig, use_container_width=True)

    # Delta summary table
    def color_delta(val):
        if isinstance(val, float):
            if val > 0.1:
                return "color: green; font-weight: bold"
            elif val < -0.1:
                return "color: red; font-weight: bold"
        return ""

    st.dataframe(
        df[["Company", "Region", "Before", "After", "Δ", "Confidence"]]
        .style.map(color_delta, subset=["Δ"]),
        hide_index=True,
        use_container_width=True,
    )
