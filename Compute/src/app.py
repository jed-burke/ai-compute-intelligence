"""
AI Compute Intelligence System — Streamlit UI
"""

import sys
import os
import io
from pathlib import Path
from datetime import datetime

# Allow imports from src/
sys.path.insert(0, str(Path(__file__).parent))

import streamlit as st
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ── Page config ───────────────────────────────────────────────────────────────
st.set_page_config(
    page_title="AI Compute Intelligence",
    page_icon="🖥️",
    layout="wide",
    initial_sidebar_state="expanded",
)

# ── Imports (after path setup) ────────────────────────────────────────────────
import ingest
import rag
import fetch
import scheduler
import dashboard
import scenario as sc


def build_export_doc(question: str, result: dict) -> bytes:
    doc = Document()

    # Title
    title = doc.add_heading("AI Compute Intelligence Report", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Date
    date_para = doc.add_paragraph(datetime.now().strftime("%B %d, %Y"))
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_para.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    doc.add_paragraph()

    # Question
    doc.add_heading("Query", level=1)
    doc.add_paragraph(question)

    # Analysis
    doc.add_heading("Analysis", level=1)
    for line in result["answer"].split("\n"):
        line = line.strip()
        if not line:
            continue
        if line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        elif line.startswith("- ") or line.startswith("* "):
            doc.add_paragraph(line[2:], style="List Bullet")
        else:
            doc.add_paragraph(line)

    # Sources
    if result["sources"]:
        doc.add_heading("Sources", level=1)
        for chunk in result["chunks"]:
            src = chunk.get("source", "")
            url = chunk.get("url", "")
            title_text = chunk.get("title", "")
            date = chunk.get("date", "")
            seen = set()
            key = url or src
            if key in seen:
                continue
            seen.add(key)
            line = f"• {src}"
            if title_text:
                line += f" — {title_text}"
            if date:
                line += f" ({date})"
            p = doc.add_paragraph(line)
            if url:
                run = p.add_run(f"\n  {url}")
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0x00, 0x56, 0xB3)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()

# ── Sidebar ───────────────────────────────────────────────────────────────────
with st.sidebar:
    st.title("🖥️ AI Compute Intelligence")
    st.caption("RAG-enabled analysis of global AI infrastructure")

    st.divider()

    # API Key
    st.subheader("🔑 Anthropic API Key")
    api_key_input = st.text_input(
        "Paste your API key here",
        type="password",
        placeholder="sk-ant-...",
        help="Get your key at console.anthropic.com",
    )
    env_key = os.getenv("ANTHROPIC_API_KEY", "")
    api_key = api_key_input or env_key

    if api_key:
        st.success("API key loaded", icon="✅")
    else:
        st.warning("No API key — add one above or in your .env file", icon="⚠️")
        st.markdown("Get a key at [console.anthropic.com](https://console.anthropic.com)")

    st.divider()

    # ── PDF Research Library ───────────────────────────────────────────────────
    st.subheader("📚 PDF Research Library")
    chunk_count = ingest.chunk_count()

    if chunk_count > 0:
        st.success(f"{chunk_count:,} chunks indexed", icon="✅")
        pdf_dir = Path(__file__).parent.parent / "research"
        pdf_count = len(list(pdf_dir.glob("*.pdf"))) if pdf_dir.exists() else 0
        if pdf_count:
            st.caption(f"{pdf_count} PDFs in corpus")
    else:
        st.info("Research library not yet indexed", icon="📂")

    if st.button("🔄 Index Research Library", use_container_width=True,
                 help="Extracts and indexes all PDFs. Takes ~2 minutes on first run."):
        progress_bar = st.progress(0, text="Starting...")
        status_text = st.empty()

        def on_progress(current, total, filename):
            pct = current / total
            short_name = filename[:45] + "..." if len(filename) > 48 else filename
            progress_bar.progress(pct, text=f"Processing {current}/{total}")
            status_text.caption(f"📄 {short_name}")

        with st.spinner("Indexing PDFs..."):
            result = ingest.ingest_all(progress_callback=on_progress)

        progress_bar.empty()
        status_text.empty()

        if result["status"] == "ok":
            st.success(
                f"Done! {result['pdfs_processed']} PDFs → "
                f"{result['new_chunks']:,} new chunks "
                f"({result['total_chunks']:,} total)",
                icon="✅",
            )
            st.rerun()
        else:
            st.error(result.get("message", "Ingestion failed"))

    st.divider()

    # ── Web Sources ────────────────────────────────────────────────────────────
    st.subheader("🌐 Web Sources")

    web_article_count = fetch.seen_url_count()
    if web_article_count > 0:
        st.success(f"{web_article_count} articles fetched", icon="✅")
    else:
        st.info("No web articles fetched yet", icon="🌐")

    sources_list = [s["name"] for s in fetch.SOURCES]
    with st.expander(f"Sources ({len(sources_list)})"):
        for name in sources_list:
            st.caption(f"• {name}")

    col_fetch, col_now = st.columns(2)
    with col_fetch:
        if st.button("🔄 Fetch Now", use_container_width=True,
                     help="Pulls new articles from all sources immediately."):
            log_area = st.empty()
            log_lines = []

            def on_fetch_progress(source_name, status):
                icon = "✅" if status.startswith("done") else ("⏳" if status == "fetching" else "⚠️")
                log_lines.append(f"{icon} **{source_name}**: {status}")
                log_area.markdown("\n\n".join(log_lines[-8:]))

            with st.spinner("Fetching web sources..."):
                articles = fetch.fetch_all(progress_callback=on_fetch_progress)

            log_area.empty()

            if articles:
                with st.spinner(f"Indexing {len(articles)} articles..."):
                    result = ingest.ingest_web_articles(articles)
                st.success(
                    f"{len(articles)} articles → {result['new_chunks']:,} new chunks",
                    icon="✅",
                )
                st.rerun()
            else:
                st.info("Already up to date.")

    st.divider()

    # ── Auto-fetch Schedule ────────────────────────────────────────────────────
    st.subheader("⏰ Auto-Fetch Schedule")

    sched_state = scheduler.get_state()
    is_enabled = sched_state.get("enabled", False)

    interval_hours = st.select_slider(
        "Fetch interval",
        options=[1, 2, 4, 6, 12, 24, 48],
        value=sched_state.get("interval_hours", 24),
        format_func=lambda h: f"Every {h}h" if h > 1 else "Every hour",
    )

    col_on, col_off = st.columns(2)
    with col_on:
        if st.button("Enable", use_container_width=True,
                     type="primary" if not is_enabled else "secondary"):
            scheduler.enable(interval_hours)
            st.rerun()
    with col_off:
        if st.button("Disable", use_container_width=True,
                     disabled=not is_enabled):
            scheduler.disable()
            st.rerun()

    if is_enabled:
        st.success("Auto-fetch active", icon="⏰")
        if sched_state.get("next_run"):
            st.caption(f"Next run: {sched_state['next_run']}")
    else:
        st.info("Auto-fetch off", icon="⏸️")

    if sched_state.get("last_run"):
        st.caption(f"Last run: {sched_state['last_run']} — {sched_state.get('last_status', '')}")

    st.divider()
    st.caption("Built with Claude · ChromaDB · Streamlit")


# ── Main area ─────────────────────────────────────────────────────────────────
st.title("AI Compute Intelligence System")

tab_rag, tab_dash, tab_scenario = st.tabs([
    "🔍 RAG Analysis", "🏢 Company Dashboard", "🔮 Scenario Forge"
])

# ════════════════════════════════════════════════════════════════════
# TAB 1 — RAG Analysis
# ════════════════════════════════════════════════════════════════════
with tab_rag:
    st.markdown(
        "Ask questions about global AI compute infrastructure, data centers, "
        "chip supply chains, export controls, and strategic dependencies — "
        "grounded in your research library and live web sources."
    )

    # Example queries
    with st.expander("💡 Example queries", expanded=chunk_count == 0):
        examples = [
            "What are the key bottlenecks in the AI chip supply chain?",
            "How are US export controls affecting global AI compute access?",
            "Which countries are building sovereign AI compute capacity?",
            "What is the role of hyperscalers in AI infrastructure?",
            "How do scaling laws relate to compute requirements for frontier models?",
            "What are the strategic dependencies in AI hardware manufacturing?",
        ]
        cols = st.columns(2)
        for i, ex in enumerate(examples):
            if cols[i % 2].button(ex, key=f"ex_{i}", use_container_width=True):
                st.session_state["query_input"] = ex
                st.rerun()

    st.divider()

    # ── Filter panel ───────────────────────────────────────────────────────────
    with st.expander("🔽 Filters", expanded=False):
        sources_info = ingest.get_sources()
        pdf_sources = sources_info["pdf"]
        web_sources = sources_info["web"]

        fcol1, fcol2, fcol3 = st.columns(3)

        with fcol1:
            st.markdown("**Source type**")
            source_type = st.radio(
                "source_type",
                options=["all", "pdf", "web"],
                format_func=lambda x: {"all": "All sources", "pdf": "PDFs only", "web": "Web only"}[x],
                label_visibility="collapsed",
            )

        with fcol2:
            st.markdown("**Specific sources**")
            if source_type == "pdf":
                available = pdf_sources
            elif source_type == "web":
                available = web_sources
            else:
                available = pdf_sources + web_sources

            selected_sources = st.multiselect(
                "specific_sources",
                options=available,
                default=[],
                placeholder="All selected" if available else "Index library first",
                label_visibility="collapsed",
            )

        with fcol3:
            st.markdown("**Date range (web articles)**")
            days_options = {
                "All time": None,
                "Last 7 days": 7,
                "Last 30 days": 30,
                "Last 90 days": 90,
                "Last 180 days": 180,
                "Last year": 365,
            }
            days_label = st.selectbox(
                "date_range",
                options=list(days_options.keys()),
                disabled=(source_type == "pdf"),
                label_visibility="collapsed",
            )
            days_back = days_options[days_label]

        where_clause = rag.build_where_clause(
            source_type=source_type,
            selected_sources=selected_sources,
            days_back=days_back,
            pdf_sources=pdf_sources,
            web_sources=web_sources,
        )

        if where_clause:
            active = []
            if source_type != "all":
                active.append(source_type.upper() + "s only")
            if selected_sources:
                active.append(f"{len(selected_sources)} source(s) selected")
            if days_back:
                active.append(f"last {days_back} days")
            st.caption("Active filters: " + " · ".join(active))

    # Query input
    query_text = st.text_area(
        "Your question",
        height=100,
        placeholder="e.g. What are the main risks to US AI compute dominance?",
        key="query_input",
    )

    col1, col2 = st.columns([1, 5])
    with col1:
        submit = st.button("Analyze", type="primary", use_container_width=True)
    with col2:
        top_k = st.slider("Sources to retrieve", min_value=4, max_value=16, value=8, step=2)

    # Guard rails
    if submit:
        if not query_text.strip():
            st.warning("Please enter a question.")
        elif not api_key:
            st.error("Add your Anthropic API key in the sidebar to run queries.")
        elif chunk_count == 0:
            st.error("Index the research library first using the sidebar button.")
        else:
            with st.spinner("Retrieving sources and analyzing..."):
                try:
                    result = rag.query(
                        query_text.strip(),
                        api_key=api_key,
                        n_results=top_k,
                        where=where_clause,
                    )
                    st.session_state["last_result"] = result
                    st.session_state["last_query"] = query_text.strip()
                except Exception as e:
                    err = str(e)
                    if "authentication" in err.lower() or "api_key" in err.lower():
                        st.error("Invalid API key. Check your key at console.anthropic.com.")
                    else:
                        st.error(f"Error: {err}")

    # Display last result
    if "last_result" in st.session_state:
        result = st.session_state["last_result"]

        col_head, col_export = st.columns([4, 1])
        with col_head:
            st.subheader("Analysis")
        with col_export:
            saved_query = st.session_state.get("last_query", query_text.strip() or "Query")
            doc_bytes = build_export_doc(saved_query, result)
            filename = f"ai_compute_report_{datetime.now().strftime('%Y%m%d_%H%M')}.docx"
            st.download_button(
                "Export .docx",
                data=doc_bytes,
                file_name=filename,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
            )

        st.markdown(result["answer"])

        if result["sources"]:
            st.divider()
            st.subheader("Sources used")
            for src in result["sources"]:
                url = next((c.get("url") for c in result["chunks"] if c.get("source") == src), None)
                if url:
                    st.markdown(f"- 🌐 [{src}]({url})")
                else:
                    st.markdown(f"- 📄 `{src}`")

        with st.expander("🔍 Retrieved chunks (debug view)"):
            for i, chunk in enumerate(result["chunks"], 1):
                url = chunk.get("url", "")
                label = f"[{chunk['source']}]({url})" if url else chunk["source"]
                st.markdown(f"**[{i}] {label}** — similarity: `{chunk['score']}`")
                st.text(chunk["text"][:400] + ("..." if len(chunk["text"]) > 400 else ""))
                st.divider()


# ════════════════════════════════════════════════════════════════════
# TAB 2 — Company Dashboard
# ════════════════════════════════════════════════════════════════════
with tab_dash:
    companies = dashboard.load_companies()

    # ── Top controls ───────────────────────────────────────────────
    dc1, dc2, dc3 = st.columns([3, 2, 1])
    with dc1:
        region_filter = st.multiselect(
            "Filter by region",
            options=["US", "PRC", "EU"],
            default=["US", "PRC", "EU"],
            format_func=lambda r: dashboard.REGION_LABELS[r],
        )
    with dc3:
        dash_report_bytes = dashboard.build_dashboard_report(companies, region_filter or ["US","PRC","EU"])
        st.download_button(
            "Export Report",
            data=dash_report_bytes,
            file_name=f"frontier_ai_report_{datetime.now().strftime('%Y%m%d')}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True,
        )

    active_regions = region_filter if region_filter else ["US", "PRC", "EU"]

    # ── Dashboard sub-tabs ─────────────────────────────────────────
    dtab_overview, dtab_map, dtab_timeline, dtab_risk, dtab_dive = st.tabs([
        "📊 Overview", "🗺️ Map", "📅 Timeline", "⚠️ Risk Matrix", "🔬 Deep Dive"
    ])

    with dtab_overview:
        dashboard.render_capability_chart(companies, active_regions)
        st.divider()
        dashboard.render_overview(companies, active_regions)

    with dtab_map:
        dashboard.render_map(companies, active_regions)

    with dtab_timeline:
        dashboard.render_timeline(companies, active_regions)

    with dtab_risk:
        rx_col, ry_col = st.columns(2)
        risk_options = list(dashboard.RISK_LABELS.keys())
        risk_x = rx_col.selectbox(
            "X axis — Risk dimension",
            options=risk_options,
            index=0,
            format_func=lambda k: dashboard.RISK_LABELS[k],
        )
        risk_y = ry_col.selectbox(
            "Y axis — Risk dimension",
            options=risk_options,
            index=1,
            format_func=lambda k: dashboard.RISK_LABELS[k],
        )
        dashboard.render_risk_matrix(companies, active_regions, risk_x, risk_y)

    with dtab_dive:
        active_companies = [c for c in companies if c["region"] in active_regions]
        company_names = [c["name"] for c in active_companies]
        selected_name = st.selectbox("Select company", options=company_names)
        selected_company = next((c for c in active_companies if c["name"] == selected_name), None)

        if selected_company:
            dashboard.render_deep_dive(selected_company)
            st.divider()

            # RAG deep-dive query
            if api_key and chunk_count > 0:
                st.markdown("**🔍 Intelligence Query — ask your corpus about this company**")
                default_q = f"What does the research say about {selected_name}'s compute strategy and supply chain dependencies?"
                dive_q = st.text_area("Query", value=default_q, height=80, key="dive_query")
                if st.button("Run Query", key="dive_submit"):
                    with st.spinner(f"Analyzing sources for {selected_name}..."):
                        dive_result = rag.query(dive_q, api_key=api_key, n_results=8)
                    st.markdown(dive_result["answer"])
                    if dive_result["sources"]:
                        st.caption("Sources: " + " · ".join(dive_result["sources"]))
            else:
                st.info("Index the research library and add an API key to run intelligence queries here.", icon="💡")


# ════════════════════════════════════════════════════════════════════
# TAB 3 — Scenario Forge
# ════════════════════════════════════════════════════════════════════
with tab_scenario:
    sc_companies = dashboard.load_companies()

    st.markdown(
        "Build a scenario, adjust the causal levers, and let the model project "
        "how the competitive landscape shifts — then ask your research corpus to explain why."
    )
    st.divider()

    # ── Preset selector ────────────────────────────────────────────
    preset_name = st.selectbox(
        "Start from a preset scenario",
        options=list(sc.PRESETS.keys()),
        help="Presets auto-populate the sliders. You can adjust any value after selecting.",
    )
    preset = sc.PRESETS[preset_name]
    st.info(preset["description"], icon="📋")

    st.divider()

    # ── Global levers ──────────────────────────────────────────────
    st.subheader("Global Variables")
    gc1, gc2, gc3, gc4 = st.columns(4)

    with gc1:
        ec_delta = st.slider(
            "Export Control Tightening",
            min_value=-5, max_value=5,
            value=preset["global"].get("export_control_delta", 0),
            help="Positive = tighter restrictions. Asymmetric impact: affects PRC most.",
        )
    with gc2:
        geo_delta = st.slider(
            "Geopolitical Tension",
            min_value=-5, max_value=5,
            value=preset["global"].get("geopolitical_tension_delta", 0),
            help="US–China tension change. Drives geopolitical exposure across all regions.",
        )
    with gc3:
        oss_delta = st.slider(
            "Open Source Momentum",
            min_value=0, max_value=10,
            value=preset["global"].get("open_source_momentum", 0),
            help="Surge in open-weight model releases. Reduces concentration risk globally.",
        )
    with gc4:
        chip_shock = st.slider(
            "Chip Supply Shock (%)",
            min_value=-60, max_value=60,
            value=preset["global"].get("chip_supply_shock", 0),
            help="Global GPU supply change (e.g. TSMC disruption). Affects all companies.",
        )

    global_params = {
        "export_control_delta": ec_delta,
        "geopolitical_tension_delta": geo_delta,
        "open_source_momentum": oss_delta,
        "chip_supply_shock": chip_shock,
    }

    # ── Company-specific overrides ─────────────────────────────────
    with st.expander("Company-Specific Overrides (optional)"):
        st.caption("Fine-tune variables for individual companies.")
        company_overrides = {}

        co_cols = st.columns(2)
        for i, company in enumerate(sc_companies):
            cid = company["id"]
            preset_co = preset["companies"].get(cid, {})
            with co_cols[i % 2]:
                with st.expander(f"{company['name']} ({company['region']})"):
                    compute_pct = st.slider(
                        "Compute Access Change (%)",
                        min_value=-80, max_value=150,
                        value=preset_co.get("compute_access_pct", 0),
                        key=f"compute_{cid}",
                    )
                    funding_bn = st.slider(
                        "Funding Change ($B)",
                        min_value=-5.0, max_value=30.0,
                        value=float(preset_co.get("funding_delta_bn", 0)),
                        step=0.5,
                        key=f"funding_{cid}",
                    )
                    model_delta = st.slider(
                        "Model Capability Jump",
                        min_value=-1.0, max_value=2.0,
                        value=float(preset_co.get("capability_model_delta", 0)),
                        step=0.1,
                        key=f"model_{cid}",
                    )
                    partnership = st.slider(
                        "New Partnership Strength",
                        min_value=0, max_value=10,
                        value=preset_co.get("partnership_strength", 0),
                        key=f"partner_{cid}",
                    )

                    if any([compute_pct, funding_bn, model_delta, partnership]):
                        company_overrides[cid] = {
                            "compute_access_pct": compute_pct,
                            "funding_delta_bn": funding_bn,
                            "capability_model_delta": model_delta,
                            "partnership_strength": partnership,
                        }

    st.divider()

    # ── Run scenario ───────────────────────────────────────────────
    run_col, _ = st.columns([1, 3])
    run_clicked = run_col.button("▶ Run Scenario", type="primary", use_container_width=True)

    if run_clicked:
        results = sc.run_scenario(
            sc_companies, preset_name, global_params, company_overrides
        )
        st.session_state["scenario_results"] = results
        st.session_state["scenario_name"] = preset_name
        st.session_state["scenario_global"] = global_params
        st.session_state["scenario_overrides"] = company_overrides
        st.session_state["scenario_narrative"] = None  # reset

    # ── Results ────────────────────────────────────────────────────
    if "scenario_results" in st.session_state:
        results = st.session_state["scenario_results"]
        s_name  = st.session_state["scenario_name"]
        s_global = st.session_state["scenario_global"]
        s_overrides = st.session_state["scenario_overrides"]

        st.divider()
        st.subheader(f"Projected Outcomes — {s_name}")

        # Summary metrics
        cap_gainers = sum(1 for r in results if r["scenario_deltas"].get("capability_score", 0) > 0.1)
        risk_increased = sum(1 for r in results if sum(r["scenario_deltas"].get(k, 0) for k in sc.RISK_KEYS) > 0.5)
        avg_conf = sum(r.get("confidence", 1) for r in results) / len(results)
        biggest_mover = max(results, key=lambda r: abs(r["scenario_deltas"].get("capability_score", 0)))

        sm1, sm2, sm3, sm4 = st.columns(4)
        sm1.metric("Companies gaining capability", cap_gainers)
        sm2.metric("Companies with elevated risk", risk_increased)
        sm3.metric("Avg model confidence", f"{int(avg_conf*100)}%")
        sm4.metric("Biggest mover", biggest_mover["name"])

        st.divider()

        # Visualizations
        vtab_cap, vtab_matrix, vtab_radar = st.tabs([
            "📊 Capability Shift", "⚠️ Risk Matrix Movement", "🕸️ Risk Radar"
        ])

        with vtab_cap:
            sc_region_filter = st.multiselect(
                "Filter regions",
                ["US", "PRC", "EU"], default=["US", "PRC", "EU"],
                key="sc_region_cap",
            )
            dashboard.render_scenario_capability_shift(results, sc_region_filter or ["US","PRC","EU"])

        with vtab_matrix:
            mx_col, my_col = st.columns(2)
            sc_risk_options = list(sc.RISK_LABELS.keys())
            sc_rx = mx_col.selectbox("X axis", sc_risk_options, index=0,
                                     format_func=lambda k: sc.RISK_LABELS[k], key="sc_rx")
            sc_ry = my_col.selectbox("Y axis", sc_risk_options, index=1,
                                     format_func=lambda k: sc.RISK_LABELS[k], key="sc_ry")
            sc_region_matrix = st.multiselect(
                "Filter regions", ["US","PRC","EU"], default=["US","PRC","EU"], key="sc_region_mx"
            )
            dashboard.render_scenario_risk_matrix(
                results, sc_region_matrix or ["US","PRC","EU"], sc_rx, sc_ry
            )

        with vtab_radar:
            radar_names = [r["name"] for r in results]
            radar_name = st.selectbox("Select company", radar_names, key="sc_radar_co")
            radar_company = next((r for r in results if r["name"] == radar_name), None)
            if radar_company:
                dashboard.render_scenario_radar(radar_company)

        st.divider()

        # ── RAG narrative ──────────────────────────────────────────
        st.subheader("📖 Intelligence Assessment")

        if not api_key:
            st.info("Add your API key in the sidebar to generate a research-grounded narrative.", icon="🔑")
        elif chunk_count == 0:
            st.info("Index the research library first to generate a grounded narrative.", icon="📚")
        else:
            if st.session_state.get("scenario_narrative"):
                st.markdown(st.session_state["scenario_narrative"])
                if st.session_state.get("scenario_sources"):
                    st.caption("Sources: " + " · ".join(st.session_state["scenario_sources"]))
            else:
                if st.button("Generate Research-Grounded Assessment", type="secondary"):
                    prompt = sc.build_scenario_prompt(s_name, s_global, s_overrides, results)
                    with st.spinner("Retrieving sources and generating assessment..."):
                        narrative_result = rag.query(prompt, api_key=api_key, n_results=10)
                    st.session_state["scenario_narrative"] = narrative_result["answer"]
                    st.session_state["scenario_sources"] = narrative_result["sources"]
                    st.rerun()

        st.divider()

        # ── Export ─────────────────────────────────────────────────
        narrative_text = st.session_state.get("scenario_narrative", "")
        report_bytes = sc.build_scenario_report(
            preset_name=s_name,
            description=preset["description"],
            global_params=s_global,
            results=results,
            narrative=narrative_text,
        )
        st.download_button(
            "📄 Export Scenario Report (.docx)",
            data=report_bytes,
            file_name=f"scenario_{s_name.replace(' ','_').lower()}_{datetime.now().strftime('%Y%m%d')}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
