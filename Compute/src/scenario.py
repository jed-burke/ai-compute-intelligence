"""
Scenario Forge — causal impact model for frontier AI scenario analysis.

Architecture:
  - Each company has baseline scores (capability + 5 risk dimensions)
  - Scenarios define changes to global and company-level variables
  - A causal weight matrix propagates those changes to scores
  - Confidence degrades as scenario magnitude grows (Bayesian prior dilution)
  - Output: before/after score sets + narrative prompt for Claude
"""

from __future__ import annotations
import copy
import io
import math
from collections import defaultdict
from datetime import datetime

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ── Score keys ─────────────────────────────────────────────────────────────────
RISK_KEYS = [
    "compute_dependency",
    "geopolitical_exposure",
    "concentration_risk",
    "regulatory_risk",
    "supply_chain_risk",
]
RISK_LABELS = {
    "compute_dependency":   "Compute Dependency",
    "geopolitical_exposure": "Geopolitical Exposure",
    "concentration_risk":   "Concentration Risk",
    "regulatory_risk":      "Regulatory Risk",
    "supply_chain_risk":    "Supply Chain Risk",
}

# ── Preset scenarios ───────────────────────────────────────────────────────────
PRESETS = {
    "Custom": {
        "description": "Build your own scenario using the sliders below.",
        "global": {
            "export_control_delta": 0,
            "geopolitical_tension_delta": 0,
            "open_source_momentum": 0,
            "chip_supply_shock": 0,
        },
        "companies": {},
    },
    "Export Controls Tightened": {
        "description": (
            "BIS/Commerce tighten H100/H200 restrictions further. "
            "New end-use checks expand to cloud providers. "
            "PRC labs face accelerated compute pressure."
        ),
        "global": {
            "export_control_delta": 3,
            "geopolitical_tension_delta": 1,
            "open_source_momentum": 0,
            "chip_supply_shock": 0,
        },
        "companies": {
            "huawei":   {"compute_access_pct": 10, "capability_model_delta": 0.3},
            "deepseek": {"compute_access_pct": -20},
            "baidu":    {"compute_access_pct": -15},
            "alibaba":  {"compute_access_pct": -15},
            "bytedance":{"compute_access_pct": -20},
        },
    },
    "PRC Achieves Compute Parity": {
        "description": (
            "Huawei Ascend 910D reaches H100-level performance. "
            "SMIC advances to 5nm yields. "
            "PRC labs regain access to large-scale training runs."
        ),
        "global": {
            "export_control_delta": 0,
            "geopolitical_tension_delta": 2,
            "open_source_momentum": 1,
            "chip_supply_shock": 0,
        },
        "companies": {
            "huawei":   {"compute_access_pct": 80, "capability_model_delta": 0.5},
            "deepseek": {"compute_access_pct": 60, "capability_model_delta": 0.8},
            "baidu":    {"compute_access_pct": 50, "capability_model_delta": 0.5},
            "alibaba":  {"compute_access_pct": 50, "capability_model_delta": 0.5},
            "bytedance":{"compute_access_pct": 40, "capability_model_delta": 0.3},
            "zhipu":    {"compute_access_pct": 30, "capability_model_delta": 0.3},
        },
    },
    "US–China Détente": {
        "description": (
            "Diplomatic reset reduces technology export friction. "
            "Some chip restrictions eased. "
            "Both sides deprioritize AI as a strategic flashpoint."
        ),
        "global": {
            "export_control_delta": -2,
            "geopolitical_tension_delta": -3,
            "open_source_momentum": 2,
            "chip_supply_shock": 0,
        },
        "companies": {
            "deepseek": {"compute_access_pct": 30},
            "baidu":    {"compute_access_pct": 20},
            "alibaba":  {"compute_access_pct": 20},
            "bytedance":{"compute_access_pct": 15},
        },
    },
    "Open Source Surge": {
        "description": (
            "Meta releases Llama 5 weights. DeepSeek and Alibaba follow. "
            "Compute-efficient open models reshape competitive dynamics. "
            "Closed-model moats narrow significantly."
        ),
        "global": {
            "export_control_delta": 0,
            "geopolitical_tension_delta": 0,
            "open_source_momentum": 6,
            "chip_supply_shock": 0,
        },
        "companies": {
            "meta_ai":  {"capability_model_delta": 1.0},
            "mistral":  {"capability_model_delta": 0.5},
            "deepseek": {"capability_model_delta": 0.8},
            "alibaba":  {"capability_model_delta": 0.6},
        },
    },
    "Global Chip Supply Crunch": {
        "description": (
            "TSMC production disruption (Taiwan Strait tension or natural disaster). "
            "NVIDIA shipments delayed 6–12 months globally. "
            "All labs face training slowdowns."
        ),
        "global": {
            "export_control_delta": 0,
            "geopolitical_tension_delta": 3,
            "open_source_momentum": 0,
            "chip_supply_shock": -40,
        },
        "companies": {},
    },
    "Frontier Capability Leap (US)": {
        "description": (
            "OpenAI or Anthropic releases a model with qualitative capability jump — "
            "strong reasoning, autonomous agents, 10x efficiency. "
            "Widens US lead and triggers regulatory response."
        ),
        "global": {
            "export_control_delta": 1,
            "geopolitical_tension_delta": 1,
            "open_source_momentum": 0,
            "chip_supply_shock": 0,
        },
        "companies": {
            "openai":    {"capability_model_delta": 1.5, "funding_delta_bn": 5},
            "anthropic": {"capability_model_delta": 0.5},
        },
    },
}


# ── Causal weight engine ───────────────────────────────────────────────────────

def _clamp(val: float, lo: float = 1.0, hi: float = 10.0) -> float:
    return max(lo, min(hi, val))


def apply_causal_model(
    company: dict,
    global_params: dict,
    company_params: dict,
) -> dict:
    """
    Compute score deltas for one company given scenario parameters.

    Global params (changes from baseline):
        export_control_delta     : -5..+5  (positive = tighter)
        geopolitical_tension_delta: -5..+5
        open_source_momentum     : 0..10
        chip_supply_shock        : -100..+100 (% change, applies to all)

    Company params (company-specific overrides):
        compute_access_pct       : % change in GPU access
        funding_delta_bn         : $B change in funding
        capability_model_delta   : direct capability score change
        partnership_strength     : 0..10 (new strategic partnerships)

    Returns dict {score_name: delta}
    """
    region = company["region"]
    deltas: dict[str, float] = defaultdict(float)

    # ── 1. Compute access (company-specific + global chip shock) ───────────────
    compute_pct = company_params.get("compute_access_pct", 0)
    chip_shock = global_params.get("chip_supply_shock", 0)
    effective_compute = compute_pct + chip_shock

    if effective_compute != 0:
        # PRC benefits more from compute gains (further from frontier)
        cap_weight = 0.025 if region == "PRC" else 0.018
        deltas["capability_score"]   += effective_compute * cap_weight
        deltas["compute_dependency"] -= effective_compute * 0.04
        deltas["supply_chain_risk"]  -= effective_compute * 0.025

    # ── 2. Funding ─────────────────────────────────────────────────────────────
    funding = company_params.get("funding_delta_bn", 0)
    if funding != 0:
        deltas["capability_score"]   += funding * 0.012
        deltas["compute_dependency"] -= funding * 0.02
        deltas["concentration_risk"] -= funding * 0.015

    # ── 3. Direct model capability delta ──────────────────────────────────────
    model_delta = company_params.get("capability_model_delta", 0)
    if model_delta != 0:
        deltas["capability_score"] += model_delta
        # More capable = more regulatory scrutiny
        deltas["regulatory_risk"]  += model_delta * 0.12

    # ── 4. Partnership strength ────────────────────────────────────────────────
    partnership = company_params.get("partnership_strength", 0)
    if partnership != 0:
        deltas["compute_dependency"] -= partnership * 0.10
        deltas["concentration_risk"] -= partnership * 0.07
        deltas["supply_chain_risk"]  -= partnership * 0.08
        deltas["capability_score"]   += partnership * 0.05

    # ── 5. Export control tightening (global, asymmetric by region) ────────────
    ec_delta = global_params.get("export_control_delta", 0)
    if ec_delta != 0:
        if region == "PRC":
            deltas["compute_dependency"]   += ec_delta * 0.22
            deltas["geopolitical_exposure"]+= ec_delta * 0.28
            deltas["supply_chain_risk"]    += ec_delta * 0.20
            deltas["capability_score"]     -= ec_delta * 0.12
            # Exception: Huawei benefits slightly from more PRC demand for Ascend
            if company.get("id") == "huawei":
                deltas["compute_dependency"]   -= ec_delta * 0.15  # partial offset
                deltas["capability_score"]     += ec_delta * 0.08
        elif region == "EU":
            deltas["compute_dependency"]   += ec_delta * 0.06
            deltas["geopolitical_exposure"]+= ec_delta * 0.10
            deltas["supply_chain_risk"]    += ec_delta * 0.05
        else:  # US
            deltas["geopolitical_exposure"]-= ec_delta * 0.04  # slightly safer
            deltas["regulatory_risk"]      += ec_delta * 0.05  # more compliance burden

    # ── 6. Geopolitical tension (global, asymmetric) ───────────────────────────
    geo_delta = global_params.get("geopolitical_tension_delta", 0)
    if geo_delta != 0:
        if region == "PRC":
            deltas["geopolitical_exposure"] += geo_delta * 0.32
            deltas["regulatory_risk"]       += geo_delta * 0.10
            deltas["concentration_risk"]    += geo_delta * 0.08
        elif region == "US":
            deltas["regulatory_risk"]       += geo_delta * 0.12
            deltas["geopolitical_exposure"] += geo_delta * 0.05
        elif region == "EU":
            deltas["regulatory_risk"]       += geo_delta * 0.18
            deltas["geopolitical_exposure"] += geo_delta * 0.12

    # ── 7. Open source momentum (global) ──────────────────────────────────────
    oss = global_params.get("open_source_momentum", 0)
    if oss != 0:
        deltas["concentration_risk"] -= oss * 0.10
        # Open-source-native companies benefit most
        if company.get("open_source"):
            deltas["capability_score"] += oss * 0.06
            deltas["concentration_risk"] -= oss * 0.05  # extra benefit

    return {k: round(v, 2) for k, v in deltas.items()}


def confidence_score(global_params: dict, company_params: dict) -> float:
    """
    Bayesian prior dilution: confidence in the projection decreases as
    scenario magnitude grows. Returns 0.0–1.0.
    """
    magnitude = (
        abs(global_params.get("export_control_delta", 0)) / 5 +
        abs(global_params.get("geopolitical_tension_delta", 0)) / 5 +
        abs(global_params.get("open_source_momentum", 0)) / 10 +
        abs(global_params.get("chip_supply_shock", 0)) / 100 +
        abs(company_params.get("compute_access_pct", 0)) / 100 +
        abs(company_params.get("capability_model_delta", 0)) / 3 +
        abs(company_params.get("funding_delta_bn", 0)) / 20 +
        abs(company_params.get("partnership_strength", 0)) / 10
    )
    # Exponential decay: confidence halves at magnitude ~1.5
    return round(math.exp(-0.45 * magnitude), 2)


def run_scenario(
    companies: list[dict],
    preset_name: str,
    global_params: dict,
    company_overrides: dict,  # {company_id: {param: value}}
) -> list[dict]:
    """
    Apply scenario to all companies. Returns list of updated company dicts
    with 'before' and 'after' scores and delta, confidence.
    """
    results = []
    for company in companies:
        cid = company["id"]
        company_params = company_overrides.get(cid, {})

        deltas = apply_causal_model(company, global_params, company_params)
        conf = confidence_score(global_params, company_params)

        before_cap = company.get("capability_score", 5)
        before_risks = dict(company.get("risks", {}))

        after_cap = _clamp(before_cap + deltas.get("capability_score", 0), 1, 10)
        after_risks = {}
        for k in RISK_KEYS:
            after_risks[k] = _clamp(
                before_risks.get(k, 5) + deltas.get(k, 0), 1, 10
            )

        results.append({
            **company,
            "scenario_before": {
                "capability_score": before_cap,
                "risks": before_risks,
            },
            "scenario_after": {
                "capability_score": round(after_cap, 2),
                "risks": {k: round(v, 2) for k, v in after_risks.items()},
            },
            "scenario_deltas": deltas,
            "confidence": conf,
            # Update live scores for downstream charts
            "capability_score": round(after_cap, 2),
            "risks": after_risks,
        })

    return results


# ── Narrative prompt builder ───────────────────────────────────────────────────

def build_scenario_prompt(
    preset_name: str,
    global_params: dict,
    company_overrides: dict,
    results: list[dict],
) -> str:
    """Build a rich RAG prompt explaining the scenario for Claude to analyze."""

    # Summarise biggest movers
    cap_movers = sorted(
        results,
        key=lambda r: abs(r["scenario_deltas"].get("capability_score", 0)),
        reverse=True,
    )[:4]

    risk_movers = sorted(
        results,
        key=lambda r: sum(abs(v) for v in r["scenario_deltas"].values()),
        reverse=True,
    )[:4]

    mover_lines = "\n".join(
        f"  - {r['name']}: capability {r['scenario_before']['capability_score']:.1f} → "
        f"{r['scenario_after']['capability_score']:.1f} "
        f"(Δ{r['scenario_deltas'].get('capability_score', 0):+.2f})"
        for r in cap_movers
    )

    risk_lines = "\n".join(
        f"  - {r['name']}: avg risk delta "
        f"{sum(r['scenario_deltas'].get(k,0) for k in RISK_KEYS)/len(RISK_KEYS):+.2f}"
        for r in risk_movers
    )

    ec = global_params.get("export_control_delta", 0)
    geo = global_params.get("geopolitical_tension_delta", 0)
    oss = global_params.get("open_source_momentum", 0)
    chip = global_params.get("chip_supply_shock", 0)

    global_summary = (
        f"Export control tightening: {ec:+.1f} | "
        f"Geopolitical tension: {geo:+.1f} | "
        f"Open-source momentum: {oss:+.1f} | "
        f"Chip supply shock: {chip:+.1f}%"
    )

    prompt = f"""I am analyzing a scenario called "{preset_name}" using a causal impact model for frontier AI competition.

SCENARIO PARAMETERS:
{global_summary}

LARGEST CAPABILITY MOVERS:
{mover_lines}

LARGEST RISK MOVERS:
{risk_lines}

Based on your research corpus, please provide a structured intelligence assessment of this scenario. Address:
1. What does the research say about the plausibility of these conditions?
2. Which companies or regions are most strategically exposed?
3. What second-order effects (supply chain, talent, regulation) are likely?
4. What are the key uncertainties that could cause the scenario to diverge from projections?
5. What signals or indicators should analysts monitor?

Be specific and cite sources where the research is relevant."""

    return prompt


# ── Scenario report export ─────────────────────────────────────────────────────

def build_scenario_report(
    preset_name: str,
    description: str,
    global_params: dict,
    results: list[dict],
    narrative: str = "",
) -> bytes:
    doc = Document()

    title = doc.add_heading("Scenario Intelligence Report", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph(f"Scenario: {preset_name}")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].bold = True
    date_p = doc.add_paragraph(datetime.now().strftime("%B %d, %Y"))
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_p.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    doc.add_paragraph()

    doc.add_heading("Scenario Description", level=1)
    doc.add_paragraph(description)

    # Parameters
    doc.add_heading("Scenario Parameters", level=1)
    param_labels = {
        "export_control_delta": "Export Control Change",
        "geopolitical_tension_delta": "Geopolitical Tension Change",
        "open_source_momentum": "Open Source Momentum",
        "chip_supply_shock": "Chip Supply Shock (%)",
    }
    for k, label in param_labels.items():
        val = global_params.get(k, 0)
        if val != 0:
            doc.add_paragraph(f"{label}: {val:+.1f}", style="List Bullet")

    # Narrative
    if narrative:
        doc.add_heading("Intelligence Assessment", level=1)
        for line in narrative.split("\n"):
            line = line.strip()
            if not line:
                continue
            if line.startswith("## "):
                doc.add_heading(line[3:], level=2)
            elif line.startswith("### "):
                doc.add_heading(line[4:], level=3)
            elif line.startswith("- ") or line.startswith("* "):
                doc.add_paragraph(line[2:], style="List Bullet")
            elif line[0].isdigit() and len(line) > 2 and line[1] in ".):":
                doc.add_paragraph(line, style="List Number")
            else:
                doc.add_paragraph(line)

    # Score changes table
    doc.add_heading("Projected Score Changes", level=1)
    cols = ["Company", "Region", "Cap Before", "Cap After", "Δ Cap", "Confidence"] + \
           [RISK_LABELS[k] + " Δ" for k in RISK_KEYS]
    table = doc.add_table(rows=1 + len(results), cols=len(cols))
    table.style = "Light List Accent 1"

    hdr = table.rows[0].cells
    for i, c in enumerate(cols):
        hdr[i].text = c

    for r_idx, company in enumerate(
        sorted(results, key=lambda x: x["scenario_deltas"].get("capability_score", 0), reverse=True), 1
    ):
        row = table.rows[r_idx].cells
        cap_delta = company["scenario_deltas"].get("capability_score", 0)
        row[0].text = company["name"]
        row[1].text = company["region"]
        row[2].text = str(company["scenario_before"]["capability_score"])
        row[3].text = str(company["scenario_after"]["capability_score"])
        row[4].text = f"{cap_delta:+.2f}"
        row[5].text = f"{int(company['confidence']*100)}%"
        for i, k in enumerate(RISK_KEYS):
            delta = company["scenario_deltas"].get(k, 0)
            row[6 + i].text = f"{delta:+.2f}" if delta != 0 else "—"

    doc.add_paragraph()
    doc.add_paragraph(
        "Note: All projections are analyst estimates based on causal impact modeling. "
        "Confidence scores reflect Bayesian prior dilution — uncertainty grows with scenario magnitude. "
        f"Report generated {datetime.now().strftime('%B %d, %Y')}."
    )

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
